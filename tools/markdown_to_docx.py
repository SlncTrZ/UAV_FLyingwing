#!/usr/bin/env python3
"""
Markdown to DOCX Converter
Chuyển đổi file markdown thành tài liệu Word với format báo cáo đẹp

Usage:
    python tools/markdown_to_docx.py <input_md> <output_docx>
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class MarkdownToDocx:
    """Chuyển đổi Markdown sang DOCX"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        self.code_block = False
    
    def setup_styles(self):
        """Thiết lập styles cho document"""
        # Default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        
        # Header styles
        for i in range(1, 4):
            heading = self.doc.styles[f'Heading {i}']
            heading.font.name = 'Times New Roman'
            heading.font.bold = True
            heading.font.size = Pt(16 - i * 2)
            heading.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
            heading.paragraph_format.space_before = Pt(12 - i * 2)
            heading.paragraph_format.space_after = Pt(6 - i * 2)
        
        # Code block style
        try:
            self.doc.styles.add_style('CodeBlock', 1)
        except:
            pass
        code_style = self.doc.styles['CodeBlock']
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9)
        code_style.font.color.rgb = RGBColor(0, 0, 0)
        code_style.paragraph_format.left_indent = Inches(0.25)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        
        # Table style
        try:
            self.doc.styles.add_style('TableHeader', 1)
        except:
            pass
        table_header = self.doc.styles['TableHeader']
        table_header.font.bold = True
        table_header.font.color.rgb = RGBColor(255, 255, 255)
        table_header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_heading(self, text, level):
        """Thêm heading"""
        # Remove # symbols
        text = re.sub(r'^#+\s*', '', text)
        self.doc.add_heading(text, level=level)
    
    def add_paragraph(self, text, style='Normal'):
        """Thêm paragraph"""
        if text.strip():
            self.doc.add_paragraph(text, style=style)
    
    def add_list_item(self, text, level=0):
        """Thêm list item"""
        paragraph = self.doc.add_paragraph(text, style='List Bullet')
        paragraph.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    
    def add_code_block(self, text):
        """Thêm code block"""
        # Remove backticks
        text = re.sub(r'^```\w*\n?', '', text)
        text = re.sub(r'\n?```$', '', text)
        text = text.strip()
        
        if text:
            paragraph = self.doc.add_paragraph(text, style='CodeBlock')
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
    
    def add_table(self, headers, rows):
        """Thêm bảng"""
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_cells[i].background_color = "4F81BD"
        
        # Data rows
        for i, row in enumerate(rows):
            row_cells = table.rows[i + 1].cells
            for j, cell in enumerate(row):
                row_cells[j].text = str(cell)
        
        self.doc.add_paragraph()
    
    def parse_table(self, lines, start_idx):
        """Parse bảng markdown"""
        headers = []
        rows = []
        
        # Parse header
        header_line = lines[start_idx]
        headers = [h.strip() for h in header_line.split('|')[1:-1]]
        
        # Parse separator
        separator = lines[start_idx + 1]
        
        # Parse rows
        idx = start_idx + 2
        while idx < len(lines) and lines[idx].startswith('|'):
            row = [cell.strip() for cell in lines[idx].split('|')[1:-1]]
            if row and any(row):
                rows.append(row)
            idx += 1
        
        return headers, rows, idx
    
    def convert(self, markdown_text):
        """Chuyển đổi markdown sang docx"""
        lines = markdown_text.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Code block
            if line.startswith('```'):
                if self.code_block:
                    # End code block
                    self.code_block = False
                else:
                    # Start code block
                    self.code_block = True
                    code_text = []
                    i += 1
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        code_text.append(lines[i])
                        i += 1
                    self.add_code_block('\n'.join(code_text))
                i += 1
                continue
            
            # Skip code block content
            if self.code_block:
                i += 1
                continue
            
            # Headers
            if line.startswith('#'):
                level = len(re.match(r'^#+', line).group())
                if level <= 3:
                    self.add_heading(line, level)
                else:
                    self.add_paragraph(line[5:].lstrip())
                i += 1
                continue
            
            # Horizontal rule
            if line.startswith('---'):
                i += 1
                continue
            
            # Table
            if line.startswith('|') and '|' in line[1:-1]:
                headers, rows, new_idx = self.parse_table(lines, i)
                if headers and rows:
                    self.add_table(headers, rows)
                    i = new_idx
                else:
                    i += 1
                continue
            
            # List item
            if line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                self.add_list_item(text)
                i += 1
                continue
            
            # Numbered list
            match = re.match(r'^\d+\.\s+', line)
            if match:
                text = line[match.end():].strip()
                paragraph = self.doc.add_paragraph(text, style='List Number')
                i += 1
                continue
            
            # Check list item with checkbox
            if line.startswith('- [') or line.startswith('* ['):
                text = line.replace('[ ]', '☐ ').replace('[x]', '☑ ')[2:].strip()
                self.add_list_item(text)
                i += 1
                continue
            
            # Regular paragraph
            self.add_paragraph(line)
            i += 1
        
        return self.doc
    
    def save(self, output_path):
        """Lưu document"""
        self.doc.save(output_path)


def main():
    if len(sys.argv) != 3:
        print("Usage: python markdown_to_docx.py <input_md> <output_docx>")
        print("Example: python markdown_to_docx.py docs/hardware/HARDWARE_WIRING_GUIDE.md docs/hardware/HARDWARE_WIRING_GUIDE.docx")
        sys.exit(1)
    
    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)
    
    print(f"Converting: {input_path}")
    print(f"Output: {output_path}")
    
    # Read markdown
    with open(input_path, 'r', encoding='utf-8') as f:
        markdown_text = f.read()
    
    # Convert to docx
    converter = MarkdownToDocx()
    doc = converter.convert(markdown_text)
    converter.save(str(output_path))
    
    print(f"[OK] Conversion complete: {output_path}")


if __name__ == '__main__':
    main()